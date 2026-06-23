from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
from scipy import stats
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\wilop\Dropbox\MPDI\2026\Auditing Explainability-corto\auditoria")
NORM_PATH = ROOT / "03_model_outputs" / "normalized" / "openai_multimodel" / "openai_multimodel_full_normalized_outputs_v1.csv"
CASE_METRICS_PATH = ROOT / "04_metrics" / "metric_tables" / "openai_multimodel" / "openai_multimodel_full_case_metrics_v1.csv"
OUT_DIR = ROOT / "04_metrics" / "inferential_analysis" / "openai_multimodel"
REPORT_DIR = ROOT / "06_reports" / "stage_reports"
SCRIPT_DIR = ROOT / "01_scripts" / "stage_04_inferential_analysis"
LOCAL_OUT = Path(r"C:\Users\wilop\Documents\Codex\2026-06-15\files-mentioned-by-the-user-proyecto\outputs")

for p in [OUT_DIR, REPORT_DIR, SCRIPT_DIR, LOCAL_OUT]:
    p.mkdir(parents=True, exist_ok=True)


METRICS = ["ECS", "EVR", "EDI"]
BOOT_N = 5000
RNG = np.random.default_rng(20260617)


def ci_bootstrap(values: np.ndarray, n: int = BOOT_N) -> tuple[float, float]:
    values = np.asarray(values, dtype=float)
    if len(values) == 0:
        return np.nan, np.nan
    idx = RNG.integers(0, len(values), size=(n, len(values)))
    means = values[idx].mean(axis=1)
    return float(np.percentile(means, 2.5)), float(np.percentile(means, 97.5))


def cliffs_delta(x: np.ndarray, y: np.ndarray) -> float:
    x = np.asarray(x, dtype=float)
    y = np.asarray(y, dtype=float)
    gt = sum(float(np.sum(xi > y)) for xi in x)
    lt = sum(float(np.sum(xi < y)) for xi in x)
    return float((gt - lt) / (len(x) * len(y)))


def cliffs_label(delta: float) -> str:
    a = abs(delta)
    if a < 0.147:
        return "negligible"
    if a < 0.33:
        return "small"
    if a < 0.474:
        return "medium"
    return "large"


def rank_biserial_from_wilcoxon(x: np.ndarray, y: np.ndarray) -> float:
    diff = np.asarray(x, dtype=float) - np.asarray(y, dtype=float)
    diff = diff[diff != 0]
    n = len(diff)
    if n == 0:
        return 0.0
    ranks = stats.rankdata(np.abs(diff))
    w_pos = float(ranks[diff > 0].sum())
    w_neg = float(ranks[diff < 0].sum())
    denom = n * (n + 1) / 2
    return float((w_pos - w_neg) / denom)


def truncate(text: str, max_chars: int = 620) -> str:
    text = " ".join(str(text).split())
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 1].rsplit(" ", 1)[0] + "..."


def zscore(series: pd.Series) -> pd.Series:
    std = series.std(ddof=0)
    if std == 0 or pd.isna(std):
        return series * 0
    return (series - series.mean()) / std


def select_representative_cases(metrics: pd.DataFrame, norm: pd.DataFrame) -> pd.DataFrame:
    df = metrics.copy()
    # High consistency: high ECS and low EVR/EDI.
    df["high_consistency_score"] = zscore(df["ECS"]) - zscore(df["EVR"]) - zscore(df["EDI"])
    # High drift: low ECS and high EVR/EDI.
    df["high_drift_score"] = -zscore(df["ECS"]) + zscore(df["EVR"]) + zscore(df["EDI"])
    center = df[METRICS].mean()
    scale = df[METRICS].std(ddof=0).replace(0, 1)
    df["center_distance"] = (((df[METRICS] - center) / scale) ** 2).sum(axis=1) ** 0.5

    chosen = [
        ("A", "High explanation consistency", df.sort_values("high_consistency_score", ascending=False).iloc[0]),
        ("B", "Intermediate consistency", df.sort_values("center_distance", ascending=True).iloc[0]),
        ("C", "High explanation drift", df.sort_values("high_drift_score", ascending=False).iloc[0]),
    ]

    rows = []
    for case_label, interpretation, row in chosen:
        subset = norm[
            (norm["case_id"] == row["case_id"])
            & (norm["model_id"] == row["model_id"])
            & (norm["model_slot"] == row["model_slot"])
        ].sort_values("repetition")
        first = subset.iloc[0]
        last = subset.iloc[-1]
        rows.append(
            {
                "case": case_label,
                "selection_type": interpretation,
                "case_id": row["case_id"],
                "task_family": row["task_family"],
                "source_dataset": row["source_dataset"],
                "model_id": row["model_id"],
                "model_slot": row["model_slot"],
                "ECS": row["ECS"],
                "EVR": row["EVR"],
                "EDI": row["EDI"],
                "prompt_excerpt": truncate(first["raw_response"] if False else get_prompt_placeholder(row["case_id"]), 400),
                "final_answer_r1": truncate(first["final_answer"], 450),
                "explanation_r1": truncate(first["explanation"], 700),
                "final_answer_r5": truncate(last["final_answer"], 450),
                "explanation_r5": truncate(last["explanation"], 700),
                "selection_score": row["high_consistency_score"] if case_label == "A" else (row["center_distance"] if case_label == "B" else row["high_drift_score"]),
            }
        )
    return pd.DataFrame(rows)


def get_prompt_placeholder(case_id: str) -> str:
    # Filled later by merge after benchmark read.
    return case_id


def descriptive_stats(metrics: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for (model_id, task_family), group in metrics.groupby(["model_id", "task_family"]):
        for metric in METRICS:
            lo, hi = ci_bootstrap(group[metric].to_numpy())
            rows.append(
                {
                    "model_id": model_id,
                    "task_family": task_family,
                    "metric": metric,
                    "n": len(group),
                    "mean": group[metric].mean(),
                    "std": group[metric].std(ddof=1),
                    "median": group[metric].median(),
                    "ci95_low": lo,
                    "ci95_high": hi,
                }
            )
    return pd.DataFrame(rows)


def paired_wilcoxon(metrics: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for task_family, group in metrics.groupby("task_family"):
        pivot_base = group.pivot(index="case_id", columns="model_id", values=METRICS)
        model_ids = sorted(group["model_id"].unique())
        if len(model_ids) != 2:
            continue
        m1, m2 = model_ids
        for metric in METRICS:
            paired = pivot_base[metric].dropna()
            if m1 not in paired.columns or m2 not in paired.columns:
                continue
            x = paired[m2].to_numpy()
            y = paired[m1].to_numpy()
            try:
                stat, p = stats.wilcoxon(x, y, zero_method="wilcox", alternative="two-sided")
            except ValueError:
                stat, p = np.nan, np.nan
            rows.append(
                {
                    "task_family": task_family,
                    "metric": metric,
                    "comparison": f"{m2} - {m1}",
                    "n_pairs": len(paired),
                    "mean_diff": float(np.mean(x - y)),
                    "median_diff": float(np.median(x - y)),
                    "wilcoxon_W": stat,
                    "p_value": p,
                    "rank_biserial": rank_biserial_from_wilcoxon(x, y),
                }
            )
    return pd.DataFrame(rows)


def mann_whitney(metrics: pd.DataFrame) -> pd.DataFrame:
    rows = []
    model_ids = sorted(metrics["model_id"].unique())
    if len(model_ids) != 2:
        return pd.DataFrame()
    m1, m2 = model_ids
    for task_family, group in metrics.groupby("task_family"):
        for metric in METRICS:
            a = group[group["model_id"] == m1][metric].to_numpy()
            b = group[group["model_id"] == m2][metric].to_numpy()
            stat, p = stats.mannwhitneyu(b, a, alternative="two-sided")
            delta = cliffs_delta(b, a)
            rows.append(
                {
                    "task_family": task_family,
                    "metric": metric,
                    "comparison": f"{m2} vs {m1}",
                    "n_model_a": len(a),
                    "n_model_b": len(b),
                    "mann_whitney_U": stat,
                    "p_value": p,
                    "cliffs_delta": delta,
                    "cliffs_label": cliffs_label(delta),
                }
            )
    return pd.DataFrame(rows)


def kruskal(metrics: pd.DataFrame) -> pd.DataFrame:
    rows = []
    for model_id, group in metrics.groupby("model_id"):
        for metric in METRICS:
            samples = [g[metric].to_numpy() for _, g in group.groupby("task_family")]
            stat, p = stats.kruskal(*samples)
            rows.append(
                {
                    "model_id": model_id,
                    "metric": metric,
                    "groups": "; ".join(sorted(group["task_family"].unique())),
                    "kruskal_H": stat,
                    "p_value": p,
                }
            )
    for metric in METRICS:
        samples = [g[metric].to_numpy() for _, g in metrics.groupby("task_family")]
        stat, p = stats.kruskal(*samples)
        rows.append(
            {
                "model_id": "all_models",
                "metric": metric,
                "groups": "; ".join(sorted(metrics["task_family"].unique())),
                "kruskal_H": stat,
                "p_value": p,
            }
        )
    return pd.DataFrame(rows)


def add_doc_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    for name, size in [("Heading 1", 15), ("Heading 2", 12), ("Heading 3", 10.5)]:
        st = doc.styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string("1F4E79")


def set_cell_text(cell, text, bold=False, size=8):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def doc_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=7.5)
        shade_cell(table.rows[0].cells[i], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=7.2)
    return table


def build_report_doc(rep: pd.DataFrame, desc: pd.DataFrame, wil: pd.DataFrame, kw: pd.DataFrame) -> Path:
    doc = Document()
    add_doc_styles(doc)
    doc.add_paragraph("Analytical Results for Representative Cases and Inferential Testing", style="Heading 1")
    p = doc.add_paragraph(
        "This document summarizes the empirical artifacts required to construct the new results subsection on representative cases of explanation consistency and drift. All analyses use the existing normalized responses and case-level metrics; no additional model calls were made."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("Representative Cases", style="Heading 2")
    rows = []
    for _, r in rep.iterrows():
        rows.append([
            r["case"],
            r["selection_type"],
            r["task_family"],
            r["model_id"],
            f"{r['ECS']:.4f}",
            f"{r['EVR']:.4f}",
            f"{r['EDI']:.4f}",
            r["explanation_r1"],
            r["explanation_r5"],
        ])
    doc_table(doc, ["Case", "Type", "Task family", "Model", "ECS", "EVR", "EDI", "Explanation r1", "Explanation r5"], rows)

    doc.add_paragraph("Key Descriptive Means and 95% CI", style="Heading 2")
    desc_pivot = desc.copy()
    rows = []
    for _, r in desc_pivot.iterrows():
        rows.append([r["model_id"], r["task_family"], r["metric"], r["n"], f"{r['mean']:.4f}", f"[{r['ci95_low']:.4f}, {r['ci95_high']:.4f}]"])
    doc_table(doc, ["Model", "Task family", "Metric", "n", "Mean", "95% CI"], rows)

    doc.add_paragraph("Wilcoxon Paired Model Comparisons", style="Heading 2")
    rows = []
    for _, r in wil.iterrows():
        rows.append([r["task_family"], r["metric"], r["comparison"], r["n_pairs"], f"{r['mean_diff']:.4f}", f"{r['p_value']:.4g}", f"{r['rank_biserial']:.4f}"])
    doc_table(doc, ["Task family", "Metric", "Comparison", "n", "Mean diff", "p", "Rank-biserial"], rows)

    doc.add_paragraph("Kruskal-Wallis Task-Family Tests", style="Heading 2")
    rows = []
    for _, r in kw.iterrows():
        rows.append([r["model_id"], r["metric"], f"{r['kruskal_H']:.4f}", f"{r['p_value']:.4g}"])
    doc_table(doc, ["Model", "Metric", "H", "p"], rows)

    out = OUT_DIR / "representative_cases_and_inferential_results_v1.docx"
    doc.save(out)
    return out


def main() -> None:
    norm = pd.read_csv(NORM_PATH)
    metrics = pd.read_csv(CASE_METRICS_PATH)
    benchmark = pd.read_csv(ROOT / "02_prompt_benchmark" / "final" / "explanation_consistency_prompt_benchmark_v1.csv")
    prompt_map = benchmark.set_index("case_id")["final_prompt"].to_dict()

    rep = select_representative_cases(metrics, norm)
    rep["prompt_text"] = rep["case_id"].map(prompt_map).map(lambda x: truncate(x, 700))
    # Put prompt before answer/explanation in final CSV.
    cols = [
        "case",
        "selection_type",
        "case_id",
        "task_family",
        "source_dataset",
        "model_id",
        "ECS",
        "EVR",
        "EDI",
        "prompt_text",
        "final_answer_r1",
        "explanation_r1",
        "final_answer_r5",
        "explanation_r5",
        "selection_score",
    ]
    rep = rep[cols]

    desc = descriptive_stats(metrics)
    wil = paired_wilcoxon(metrics)
    mw = mann_whitney(metrics)
    kw = kruskal(metrics)

    rep_path = OUT_DIR / "representative_cases_consistency_drift_v1.csv"
    desc_path = OUT_DIR / "descriptive_stats_ci95_v1.csv"
    wil_path = OUT_DIR / "wilcoxon_model_comparisons_v1.csv"
    mw_path = OUT_DIR / "mannwhitney_cliffs_delta_v1.csv"
    kw_path = OUT_DIR / "kruskal_task_family_tests_v1.csv"
    rep.to_csv(rep_path, index=False, encoding="utf-8-sig")
    desc.to_csv(desc_path, index=False, encoding="utf-8-sig")
    wil.to_csv(wil_path, index=False, encoding="utf-8-sig")
    mw.to_csv(mw_path, index=False, encoding="utf-8-sig")
    kw.to_csv(kw_path, index=False, encoding="utf-8-sig")

    # Markdown summary for manuscript drafting.
    md = ["# Representative Cases and Inferential Analysis", ""]
    md.append("## Representative cases")
    md.append("")
    md.append("| Case | Type | Task family | Model | ECS | EVR | EDI |")
    md.append("| --- | --- | --- | --- | ---: | ---: | ---: |")
    for _, r in rep.iterrows():
        md.append(f"| {r['case']} | {r['selection_type']} | {r['task_family']} | {r['model_id']} | {r['ECS']:.4f} | {r['EVR']:.4f} | {r['EDI']:.4f} |")
    md.append("")
    md.append("## Significant paired model comparisons (Wilcoxon p < .05)")
    sig_wil = wil[wil["p_value"] < 0.05].copy()
    if len(sig_wil) == 0:
        md.append("No paired model comparison reached p < .05.")
    else:
        md.append("| Task family | Metric | Mean diff | p | Rank-biserial |")
        md.append("| --- | --- | ---: | ---: | ---: |")
        for _, r in sig_wil.iterrows():
            md.append(f"| {r['task_family']} | {r['metric']} | {r['mean_diff']:.4f} | {r['p_value']:.4g} | {r['rank_biserial']:.4f} |")
    md.append("")
    md.append("## Kruskal-Wallis task-family tests")
    md.append("| Model | Metric | H | p |")
    md.append("| --- | --- | ---: | ---: |")
    for _, r in kw.iterrows():
        md.append(f"| {r['model_id']} | {r['metric']} | {r['kruskal_H']:.4f} | {r['p_value']:.4g} |")

    report_path = REPORT_DIR / "stage_06_representative_cases_and_inferential_analysis.md"
    report_path.write_text("\n".join(md), encoding="utf-8")

    docx_path = build_report_doc(rep, desc, wil, kw)
    script_dst = SCRIPT_DIR / "run_representative_cases_and_stats.py"
    script_dst.write_text(Path(__file__).read_text(encoding="utf-8"), encoding="utf-8")

    manifest = {
        "stage": "06_representative_cases_and_inferential_analysis",
        "inputs": [str(NORM_PATH), str(CASE_METRICS_PATH)],
        "outputs": [str(rep_path), str(desc_path), str(wil_path), str(mw_path), str(kw_path), str(report_path), str(docx_path), str(script_dst)],
        "bootstrap_iterations": BOOT_N,
        "representative_cases": rep[["case", "case_id", "task_family", "model_id", "ECS", "EVR", "EDI"]].to_dict("records"),
    }
    manifest_path = OUT_DIR / "representative_cases_and_stats_manifest_v1.json"
    manifest_path.write_text(json.dumps(manifest, indent=2, ensure_ascii=False), encoding="utf-8")

    for output in manifest["outputs"] + [str(manifest_path)]:
        print(output)
    # Copy user-facing report artifacts.
    for path in [rep_path, desc_path, wil_path, mw_path, kw_path, report_path, docx_path]:
        target = LOCAL_OUT / Path(path).name
        target.write_bytes(Path(path).read_bytes())


if __name__ == "__main__":
    main()
